# Scan the entire DB for upcoming events (not just log.json)
def fetch_all_upcoming_events(limit=30):
    import os, json, datetime
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    now = datetime.datetime.now()
    events = []
    # Candidates: interviews, pending approvals, onboarding
    candidate_file = os.path.join(db_folder, 'candidates.json')
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r', encoding='utf-8') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
        for c in candidates:
            # Upcoming interview (check with or without time)
            date_str = c.get('interview_date', '')
            time_str = c.get('interview_time', '')
            if date_str:
                try:
                    if time_str:
                        dt = datetime.datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M")
                    else:
                        dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
                    if dt >= now:
                        events.append({
                            'type': 'Interview',
                            'date': date_str,
                            'time': time_str,
                            'name': c.get('name', ''),
                            'job_title': c.get('job_title', c.get('position', '')),
                            'candidate_id': c.get('id', '')
                        })
                except Exception:
                    pass
            # Pending approval
            if c.get('status') in ['Pending Approval', 'Selected']:
                applied_date = c.get('applied_date', '')
                events.append({
                    'type': 'Approval',
                    'date': applied_date,
                    'time': '',
                    'name': c.get('name', ''),
                    'job_title': c.get('job_title', c.get('position', '')),
                    'candidate_id': c.get('id', '')
                })
            # Onboarding
            onboarding = c.get('onboarding')
            if c.get('status') == 'Hired' and onboarding and isinstance(onboarding, dict):
                onboarding_start = onboarding.get('start_date', '')
                if onboarding_start:
                    try:
                        dt = datetime.datetime.strptime(onboarding_start, "%Y-%m-%d")
                        if dt >= now:
                            events.append({
                                'type': 'Onboarding',
                                'date': onboarding_start,
                                'time': '',
                                'name': c.get('name', ''),
                                'job_title': c.get('job_title', c.get('position', '')),
                                'candidate_id': c.get('id', '')
                            })
                    except Exception:
                        pass
    # Jobs: future job postings (if any with a future posted_at)
    job_file = os.path.join(db_folder, 'jobs.json')
    if os.path.exists(job_file):
        with open(job_file, 'r', encoding='utf-8') as f:
            try:
                jobs = json.load(f)
            except json.JSONDecodeError:
                jobs = []
        for job in jobs:
            posted_at = job.get('posted_at', '')
            if posted_at:
                try:
                    dt = datetime.datetime.strptime(posted_at, "%Y-%m-%d %H:%M:%S")
                except Exception:
                    try:
                        dt = datetime.datetime.strptime(posted_at, "%Y-%m-%d")
                    except Exception:
                        continue
                if dt >= now:
                    events.append({
                        'type': 'Job',
                        'date': posted_at.split(' ')[0],
                        'time': posted_at.split(' ')[1] if ' ' in posted_at else '',
                        'name': job.get('job_title', ''),
                        'job_title': job.get('job_title', ''),
                        'job_id': job.get('job_id', job.get('id', ''))
                    })
    # Notifications: future-dated notifications
    notification_file = os.path.join(db_folder, 'notifications.json')
    if os.path.exists(notification_file):
        with open(notification_file, 'r', encoding='utf-8') as f:
            try:
                notifications = json.load(f)
            except json.JSONDecodeError:
                notifications = []
        for n in notifications:
            date_str = n.get('date', '')
            time_str = n.get('time', '')
            if date_str:
                try:
                    dt = datetime.datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M") if time_str else datetime.datetime.strptime(date_str, "%Y-%m-%d")
                    if dt >= now:
                        events.append({
                            'type': 'Notification',
                            'date': date_str,
                            'time': time_str,
                            'name': n.get('title', ''),
                            'description': n.get('description', ''),
                        })
                except Exception:
                    pass
    # Sort by date/time ascending (soonest first)
    def event_sort_key(e):
        try:
            return datetime.datetime.strptime(e['date'] + (f" {e['time']}" if e['time'] else ''), "%Y-%m-%d %H:%M")
        except Exception:
            try:
                return datetime.datetime.strptime(e['date'], "%Y-%m-%d")
            except Exception:
                return datetime.datetime.max
    events = [e for e in events if e['date']]
    events = sorted(events, key=event_sort_key)
    return events[:limit]
import threading

# Log an event/action to db/log.json
def log_event(action_type, description, user, related_id=None, date=None, time=None, extra=None):
    import os, json, datetime
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    log_file = os.path.join(db_folder, 'log.json')
    log_entry = {
        'timestamp': date if date else datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'type': action_type,  # e.g., 'interview', 'approval', 'job', 'onboarding', etc.
        'description': description,
        'user': user,
        'related_id': related_id,
        'time': time,
        'extra': extra or {}
    }
    def write_log():
        try:
            if os.path.exists(log_file):
                with open(log_file, 'r', encoding='utf-8') as f:
                    logs = json.load(f)
            else:
                logs = []
            logs.append(log_entry)
            with open(log_file, 'w', encoding='utf-8') as f:
                json.dump(logs, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error logging event: {e}")
    # Write asynchronously to avoid blocking
    threading.Thread(target=write_log).start()

# Fetch events from log.json (optionally filter by type, date, etc.)
def fetch_events_from_log(event_types=None, upcoming_only=True, limit=20):
    import os, json, datetime
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    log_file = os.path.join(db_folder, 'log.json')
    now = datetime.datetime.now()
    events = []
    if os.path.exists(log_file):
        with open(log_file, 'r', encoding='utf-8') as f:
            try:
                logs = json.load(f)
            except json.JSONDecodeError:
                logs = []
        for entry in logs:
            if event_types and entry.get('type') not in event_types:
                continue
            # For upcoming, check if date/time is in the future
            if upcoming_only:
                date_str = entry.get('timestamp', '')
                time_str = entry.get('time', '')
                try:
                    dt = datetime.datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
                except Exception:
                    try:
                        dt = datetime.datetime.strptime(date_str, '%Y-%m-%d')
                    except Exception:
                        continue
                if dt < now:
                    continue
            events.append(entry)
    # Sort by date/time ascending (soonest first)
    events = sorted(events, key=lambda x: (x.get('timestamp', ''), x.get('time', '') or ''))
    return events[:limit]
# Utility to fetch all relevant DB data for AI context
def fetch_all_db_data():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    # Candidates
    candidate_file = os.path.join(db_folder, 'candidates.json')
    candidates = []
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
    # Jobs
    job_file = os.path.join(db_folder, 'jobs.json')
    jobs = []
    if os.path.exists(job_file):
        with open(job_file, 'r') as f:
            try:
                jobs = json.load(f)
            except json.JSONDecodeError:
                jobs = []
    # Activities
    from data import fetch_recent_activities
    activities = fetch_recent_activities(show_all=True)
    # User data
    user_file = os.path.join(db_folder, 'userdata.json')
    users = []
    if os.path.exists(user_file):
        with open(user_file, 'r') as f:
            try:
                users = json.load(f)
            except json.JSONDecodeError:
                users = []
    # Notifications
    notification_file = os.path.join(db_folder, 'notifications.json')
    notifications = []
    if os.path.exists(notification_file):
        with open(notification_file, 'r') as f:
            try:
                notifications = json.load(f)
            except json.JSONDecodeError:
                notifications = []
    # Return all as dict
    return {
        "candidates": candidates,
        "jobs": jobs,
        "activities": activities,
        "users": users,
        "notifications": notifications
    }
# Enhanced Recent Activities fetcher with new activity logger integration
def fetch_recent_activities(show_all=False):
    """
    Fetch recent activities using both legacy data and new activity logger.
    Returns combined activities from multiple sources for comprehensive view.
    """
    all_activities = []
    
    # Try to get activities from new activity logger first
    try:
        from activity_logger import activity_logger
        new_activities = activity_logger.get_recent_activities(limit=20, days=7)
        
        # Convert new activity format to expected format
        for activity in new_activities:
            all_activities.append({
                'date': activity.get('timestamp', activity.get('date', '')),
                'description': activity.get('description', ''),
                'user': activity.get('user', 'Unknown'),
                'is_new': True,
                'type': activity.get('activity_type', 'unknown'),
                'entity_type': activity.get('entity_type', ''),
                'entity_id': activity.get('entity_id', '')
            })
        
        print(f"✅ Fetched {len(new_activities)} activities from new logger")
    except Exception as e:
        print(f"⚠️ Could not fetch from new activity logger: {e}")
    
    # Fallback to legacy activity fetching from JSON files
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    job_file = os.path.join(db_folder, 'jobs.json')
    
    # Get username mapping
    usernames = {}
    user_file = os.path.join(db_folder, 'userdata.json')
    if os.path.exists(user_file):
        with open(user_file, 'r') as f:
            try:
                users = json.load(f)
            except json.JSONDecodeError:
                users = []
        for u in users:
            if u.get('email'):
                usernames[u['email']] = u.get('username', '')
            if u.get('username'):
                usernames[u['username']] = u.get('username', '')
    
    # Legacy candidate activities
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
        
        for c in candidates:
            # Initial application
            if c.get('applied_date'):
                user_val = c.get('email', '')
                user_val = usernames.get(user_val, user_val)
                all_activities.append({
                    'date': c.get('applied_date'),
                    'description': f"Candidate {c.get('name', '')} (ID: {c.get('id', '')}) applied",
                    'user': user_val if user_val else 'Unknown',
                    'is_new': True,
                    'type': 'candidate_applied',
                    'entity_type': 'candidate',
                    'entity_id': c.get('id', '')
                })
            
            # Updates
            if c.get('updated_at') and c.get('updated_by'):
                user_val = c.get('updated_by', c.get('email', ''))
                user_val = usernames.get(user_val, user_val)
                date_val = c.get('updated_at')
                candidate_name = c.get('name', '')
                candidate_id = c.get('id', '')
                
                # Status changes
                if c.get('status'):
                    all_activities.append({
                        'date': date_val,
                        'description': f"Candidate {candidate_name} (ID: {candidate_id}) status changed to {c.get('status', '')}",
                        'user': user_val if user_val else 'Unknown',
                        'is_new': True,
                        'type': 'candidate_status_updated',
                        'entity_type': 'candidate',
                        'entity_id': candidate_id
                    })
                
                # Interview scheduling
                if c.get('interview_date') and c.get('interview_time'):
                    all_activities.append({
                        'date': date_val,
                        'description': f"Interview scheduled for {candidate_name} (ID: {candidate_id}) on {c.get('interview_date', '')} at {c.get('interview_time', '')}",
                        'user': user_val if user_val else 'Unknown',
                        'is_new': True,
                        'type': 'interview_scheduled',
                        'entity_type': 'interview',
                        'entity_id': candidate_id
                    })
                
                # Onboarding updates
                if c.get('onboarding'):
                    all_activities.append({
                        'date': date_val,
                        'description': f"Onboarding updated for {candidate_name} (ID: {candidate_id})",
                        'user': user_val if user_val else 'Unknown',
                        'is_new': True,
                        'type': 'onboarding_updated',
                        'entity_type': 'onboarding',
                        'entity_id': candidate_id
                    })
    
    # Legacy job activities
    if os.path.exists(job_file):
        with open(job_file, 'r') as f:
            try:
                jobs = json.load(f)
            except json.JSONDecodeError:
                jobs = []
        
        for job in jobs:
            user_val = job.get('job_posted_by', '')
            user_val = usernames.get(user_val, user_val)
            all_activities.append({
                'date': job.get('posted_at', ''),
                'description': f"Job posted: {job.get('job_title', '')}",
                'user': user_val,
                'is_new': False,
                'type': 'job_posted',
                'entity_type': 'job',
                'entity_id': job.get('id', '')
            })
    
    # Remove duplicates based on description and date
    unique_activities = []
    seen = set()
    for activity in all_activities:
        key = f"{activity.get('date', '')}_{activity.get('description', '')}"
        if key not in seen:
            seen.add(key)
            unique_activities.append(activity)
    
    # Sort by date descending
    try:
        unique_activities = sorted(unique_activities, key=lambda x: x.get('date', ''), reverse=True)
    except Exception as e:
        print(f"⚠️ Error sorting activities: {e}")
    
    # Mark activities as new if they're within the last 2 days
    for a in unique_activities:
        try:
            date_str = a.get('date', '')
            if not date_str:
                a['is_new'] = False
                continue
                
            # Try different date formats
            dt = None
            for fmt in ['%Y-%m-%d %H:%M:%S', '%Y-%m-%d']:
                try:
                    dt = datetime.datetime.strptime(date_str, fmt)
                    break
                except ValueError:
                    continue
            
            if dt:
                a['is_new'] = (datetime.datetime.now() - dt).days < 2
            else:
                a['is_new'] = False
        except Exception:
            a['is_new'] = False
    
    # Return the most recent activities
    result = unique_activities[:20] if show_all else unique_activities[:10]
    print(f"✅ Returning {len(result)} total activities")
    return result

# ------------------------------------------------------------------------------------
# Extract text from uploaded JD file (txt, pdf, docx)

import re
from threading import Thread
from werkzeug.utils import secure_filename
import datetime
import os
import json
import openai


def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == ".pdf":
            import PyPDF2
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
    except Exception as e:
        text = ""
    return text

# Use OpenAI to extract JD from text
def extract_jd_with_openai(jd_text):
    import openai
    try:
        messages = [
            {
                "role": "system",
                "content": "You are an expert HR assistant. Extract and return only the job description from the following text. Do not include any explanations or comments."
            },
            {
                "role": "user",
                "content": f"Extract the job description from this text:\n{jd_text}"
            }
        ]
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=messages,
            temperature=0
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        return ""

def extract_audio_from_video(video_path, audio_path):
    """Extract audio from video using moviepy."""
    try:
        from moviepy import VideoFileClip
        video = VideoFileClip(video_path)
        if not video.audio:
            print("No audio stream found in the video.")
            return False
        video.audio.write_audiofile(audio_path, fps=16000, nbytes=2, codec='pcm_s16le')
        video.close()
        return True
    except ImportError:
        print("moviepy is not installed. Install it with: pip install moviepy")
        return False
    except Exception as e:
        print(f"Audio extraction error (moviepy): {e}")
        return False

def transcribe_with_whisper(audio_path):
    try:
        with open(audio_path, "rb") as audio_file:
            transcript = openai.Audio.transcribe(
                model="whisper-1",
                file=audio_file,
                response_format="text"
            )
        return transcript
    except Exception as e:
        print(f"OpenAI Whisper API transcription error: {e}")
        return None

def analyze_transcript_with_openai(transcript):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """You are an expert interviewer analyzing candidate performance.\n\nAnalyze the interview transcript and provide:\n1. A comprehensive performance summary\n2. Specific feedback on communication skills, technical knowledge, and overall interview performance\n3. Areas for improvement\n4. A final performance score out of 100\n\nIMPORTANT: Always end your response with a clear score in this exact format: \"Performance Score: X/100\" where X is a number between 0-100."""},
                {"role": "user", "content": f"Please analyze this interview transcript and provide detailed feedback with a performance score:\n\n{transcript}"}
            ],
            temperature=0.0,
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        print(f"OpenAI summary error: {e}")
        return None


def analyze_cv_with_jd(cv_data, jd_text):
    """
    Simple keyword-based match score between candidate CV and job description.
    Returns an integer score (0-100).
    """
    if not jd_text or not cv_data:
        return 0
    # Extract keywords from JD
    import re
    jd_keywords = set(re.findall(r'\b\w+\b', jd_text.lower()))
    # Extract candidate skills and experience as keywords
    cv_keywords = set()
    for field in ['skills', 'experience', 'education', 'certifications', 'projects', 'position']:
        value = cv_data.get(field, '')
        if isinstance(value, list):
            for item in value:
                cv_keywords.update(re.findall(r'\b\w+\b', str(item).lower()))
        elif isinstance(value, str):
            cv_keywords.update(re.findall(r'\b\w+\b', value.lower()))
    # Calculate intersection
    match_count = len(jd_keywords & cv_keywords)
    total_keywords = len(jd_keywords) if jd_keywords else 1
    score = int((match_count / total_keywords) * 100)
    return score

def extract_score_from_summary(summary):
    """
    Extract performance score from AI analysis summary.
    Handles multiple score formats that AI might return.
    """
    if not summary:
        return None
    summary_text = str(summary)
    patterns = [
        r'[Ss]core[:\s]*(\d{1,3})\s*/\s*100',
        r'[Ss]core[:\s]*(\d{1,3})\s*out\s*of\s*100',
        r'[Pp]erformance[^:]*[:\s]*(\d{1,3})\s*/\s*100',
        r'[Pp]erformance[^:]*[:\s]*(\d{1,3})\s*out\s*of\s*100',
        r'[Rr]ating[^:]*[:\s]*(\d{1,3})\s*/\s*100',
        r'[Rr]ating[^:]*[:\s]*(\d{1,3})\s*out\s*of\s*100',
        r'(\d{1,3})\s*/\s*100',
        r'(\d{1,3})\s*out\s*of\s*100',
        r'[Ss]core\s*of\s*(\d{1,3})',
        r'[Pp]erformance\s*of\s*(\d{1,3})',
        r'(\d{1,3})%',
        r'(\d{1,3})\s*points',
        r'scored\s*(\d{1,3})',
        r'(\d{1,3}).*(?:score|rating|performance|points)',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, summary_text, re.IGNORECASE)
        for match in matches:
            try:
                score = int(match)
                if 0 <= score <= 100:
                    print(f"[DEBUG] Extracted score {score} using pattern: {pattern}")
                    return score
            except (ValueError, TypeError):
                continue
    print(f"[WARNING] Could not extract score from summary: {summary_text[:200]}...")
    return None




# ------------------------------------------------------------------------------------
def fetch_user_data():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    user_file = os.path.join(db_folder, 'userdata.json')
    user_data = []
    if os.path.exists(user_file):
        with open(user_file, 'r') as f:
            try:
                user_data = json.load(f)
            except json.JSONDecodeError:
                user_data = []
    return user_data

def total_users():
    user_data = fetch_user_data()
    return len(user_data)

def edit_user_data(username, new_data):
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    user_file = os.path.join(db_folder, 'userdata.json')
    user_data = fetch_user_data()
    updated = False
    for user in user_data:
        if user.get('username') == username:
            user.update(new_data)
            updated = True
            break
    if updated:
        with open(user_file, 'w') as f:
            json.dump(user_data, f, indent=4)
    return updated



# ------------------------------------------------------------------------------------

def fetch_job_data():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    job_file = os.path.join(db_folder, 'jobs.json')
    job_data = []
    if os.path.exists(job_file):
        with open(job_file, 'r') as f:
            try:
                job_data = json.load(f)
            except json.JSONDecodeError:
                job_data = []
    return job_data

def job_count():
    return len(fetch_job_data())

def openings_count():
    job_data = fetch_job_data()
    openings_count = 0
    for job in job_data:
        # Only count jobs that are "Open" or don't have a status (default to Open)
        job_status = job.get('status', 'Open').lower()
        if job_status == 'open':
            # Try common key names for openings
            for key in ['job_openings', 'openings', 'openings_count', 'vacancies']:
                if key in job:
                    try:
                        openings_count += int(job[key])
                        break  # Only count one key per job
                    except (ValueError, TypeError):
                        continue
    return openings_count

def open_vacancies_count():
    """Count vacancies for open jobs only"""
    job_data = fetch_job_data()
    # We need to import the function from app.py or define it here
    # For now, let's use a simple logic similar to calculate_automatic_job_status
    from datetime import datetime, timedelta
    candidates = fetch_candidate_data()
    
    open_count = 0
    current_date = datetime.now()
    
    for job in job_data:
        # Calculate if job should be considered open
        is_open = True
        
        # Check lead time expiration
        try:
            posted_date = datetime.strptime(job.get('posted_at', ''), '%Y-%m-%d %H:%M:%S')
            lead_time_days = int(job.get('job_lead_time', 30))
            expiration_date = posted_date + timedelta(days=lead_time_days)
            if current_date > expiration_date:
                is_open = False
        except (ValueError, TypeError):
            pass  # If we can't parse dates, assume not expired
        
        # Check if all positions are filled
        if is_open:
            job_id_str = str(job.get('job_id', ''))
            hired_count = sum(1 for c in candidates 
                             if c.get('status', '').lower() == 'hired' and 
                                str(c.get('job_id', '')) == job_id_str)
            try:
                job_openings = int(job.get('job_openings', 0))
                if hired_count >= job_openings:
                    is_open = False
            except (ValueError, TypeError):
                pass
        
        # Check explicit status override
        job_status = job.get('status', '').lower()
        if job_status in ['closed', 'filled', 'cancelled', 'expired', 'on hold']:
            is_open = False
        
        if is_open:
            for key in ['job_openings', 'openings', 'openings_count', 'vacancies']:
                if key in job:
                    try:
                        open_count += int(job[key])
                        break
                    except (ValueError, TypeError):
                        continue
    return open_count

def closed_vacancies_count():
    """Count vacancies for closed jobs only"""
    job_data = fetch_job_data()
    from datetime import datetime, timedelta
    candidates = fetch_candidate_data()
    
    closed_count = 0
    current_date = datetime.now()
    
    for job in job_data:
        # Calculate if job should be considered closed
        is_closed = False
        
        # Check lead time expiration
        try:
            posted_date = datetime.strptime(job.get('posted_at', ''), '%Y-%m-%d %H:%M:%S')
            lead_time_days = int(job.get('job_lead_time', 30))
            expiration_date = posted_date + timedelta(days=lead_time_days)
            if current_date > expiration_date:
                is_closed = True
        except (ValueError, TypeError):
            pass  # If we can't parse dates, assume not expired
        
        # Check if all positions are filled
        if not is_closed:
            job_id_str = str(job.get('job_id', ''))
            hired_count = sum(1 for c in candidates 
                             if c.get('status', '').lower() == 'hired' and 
                                str(c.get('job_id', '')) == job_id_str)
            try:
                job_openings = int(job.get('job_openings', 0))
                if hired_count >= job_openings:
                    is_closed = True
            except (ValueError, TypeError):
                pass
        
        # Check explicit status override
        job_status = job.get('status', '').lower()
        if job_status in ['closed', 'filled', 'cancelled', 'expired', 'on hold']:
            is_closed = True
        
        if is_closed:
            for key in ['job_openings', 'openings', 'openings_count', 'vacancies']:
                if key in job:
                    try:
                        closed_count += int(job[key])
                        break
                    except (ValueError, TypeError):
                        continue
    return closed_count

def no_status_vacancies_count():
    """Count vacancies for jobs without status"""
    job_data = fetch_job_data()
    no_status_count = 0
    for job in job_data:
        job_status = job.get('status', '')
        if not job_status:  # No status field or empty status
            for key in ['job_openings', 'openings', 'openings_count', 'vacancies']:
                if key in job:
                    try:
                        no_status_count += int(job[key])
                        break
                    except (ValueError, TypeError):
                        continue
    return no_status_count

def total_all_vacancies_count():
    """Count ALL vacancies (open + closed + no status)"""
    job_data = fetch_job_data()
    total_count = 0
    for job in job_data:
        for key in ['job_openings', 'openings', 'openings_count', 'vacancies']:
            if key in job:
                try:
                    total_count += int(job[key])
                    break
                except (ValueError, TypeError):
                    continue
    return total_count

    
# ------------------------------------------------------------------------------------


def fetch_candidate_data():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    candidate_data = []
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidate_data = json.load(f)
            except json.JSONDecodeError:
                candidate_data = []
    return candidate_data

def candidate_count():
    candidate_data = fetch_candidate_data()
    return len(candidate_data)

# Event functions for dashboard
def fetch_upcoming_interviews():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    upcoming = []
    now = datetime.datetime.now()
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
        for c in candidates:
            date_str = c.get('interview_date', '')
            time_str = c.get('interview_time', '')
            if date_str and time_str:
                try:
                    interview_dt = datetime.datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M")
                    if interview_dt >= now:
                        upcoming.append(c)
                except Exception:
                    continue
    return upcoming

def fetch_pending_approvals_events():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    pending = []
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
        for c in candidates:
            if c.get('status') == 'Pending Approval':
                pending.append(c)
    return pending

def fetch_onboarding_events():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    onboarding = []
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
        for c in candidates:
            if c.get('status') == 'Hired' and 'onboarding' in c:
                onboarding.append(c)
    return onboarding

def edit_candidate_data(candidate_id, new_data):
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    candidate_data = fetch_candidate_data()
    updated = False
    for candidate in candidate_data:
        if candidate.get('id') == candidate_id:
            candidate.update(new_data)
            updated = True
            break
    if updated:
        with open(candidate_file, 'w') as f:
            json.dump(candidate_data, f, indent=4)
    return updated

def fetch_candidates_by_filter(**filters):
    """
    Fetch candidates matching all provided filter key-value pairs.

    Example:
        fetch_candidates_by_filter(status='active', department='HR')
    """
    candidate_data = fetch_candidate_data()
    filtered_candidates = []
    for candidate in candidate_data:
        if all(candidate.get(key) == value for key, value in filters.items()):
            filtered_candidates.append(candidate)
    return filtered_candidates

def fetch_pending_approvals():
    """
    Return a list of dicts with candidate id and name whose status is 'Selected' (pending approval).
    """
    candidate_data = fetch_candidate_data()
    pending_approvals = [
        {"id": candidate.get("id"), "name": candidate.get("name")}
        for candidate in candidate_data
        if candidate.get("status") == "Selected"
    ]
    return pending_approvals


def get_sender():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    candidate_data = fetch_candidate_data()
    intervier = None
    for candidate in candidate_data:
        if candidate.get('status') == 'Selected':
            intervier = candidate.get('intervier', 'Unknown')
    if intervier:
        return intervier
    else:
        return "Unknown"



# ------------------------------------------------------------------------------------

def fetch_onboarding_data():
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    onboarding_file = os.path.join(db_folder, 'onboarding.json')
    onboarding_data = []
    if os.path.exists(onboarding_file):
        with open(onboarding_file, 'r') as f:
            try:
                onboarding_data = json.load(f)
            except json.JSONDecodeError:
                onboarding_data = []
    return onboarding_data




# ------------------------------------------------------------------------------------
def extract_resume_with_openai(resume_path):
    """Extract structured candidate data from a resume using OpenAI."""
    if not os.path.exists(resume_path):
        return {"error": "Resume file not found."}
    
    try:
        ext = os.path.splitext(resume_path)[1].lower()
        resume_text = ""
        
        # Extract text from file
        if ext == ".pdf":
            try:
                import PyPDF2
                with open(resume_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        resume_text += page.extract_text() or ""
            except ImportError:
                return {"error": "PyPDF2 not installed for PDF processing"}
            except Exception as pdf_error:
                return {"error": f"PDF processing failed: {str(pdf_error)}"}
                
        elif ext in [".docx"]:
            try:
                import docx
                doc = docx.Document(resume_path)
                resume_text = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                return {"error": "python-docx not installed for DOCX processing"}
            except Exception as docx_error:
                return {"error": f"DOCX processing failed: {str(docx_error)}"}
        else:
            try:
                with open(resume_path, "r", encoding="utf-8") as file:
                    resume_text = file.read()
            except UnicodeDecodeError:
                try:
                    with open(resume_path, "r", encoding="latin-1") as file:
                        resume_text = file.read()
                except Exception as read_error:
                    return {"error": f"File reading failed: {str(read_error)}"}

        if not resume_text.strip():
            return {"error": "No text could be extracted from the resume"}

        # Check if OpenAI API key is configured
        if not openai.api_key:
            return {"error": "OpenAI API key not configured"}

        # Prepare OpenAI request with error handling
        messages = [
            {
                "role": "system",
                "content": (
                    "You extract structured data from resumes. "
                    "Return only a JSON object with no explanations or comments."
                )
            },
            {
                "role": "user",
                "content": (
                    f"Extract the following fields from this resume and return JSON only:\n"
                    f"name, email, phone, skills (as list), experience (as list), education (as list), "
                    f"certifications (as list), projects (as list), linkedin, github, job_title.\n\n"
                    f"Resume:\n{resume_text[:4000]}"  # Limit text to avoid token limits
                )
            }
        ]

        # Use try-except specifically for OpenAI API call
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4o",
                messages=messages,
                temperature=0
            )
            return response['choices'][0]['message']['content']
            
        except ImportError as import_error:
            return {"error": f"OpenAI library import failed: {str(import_error)}"}
        except AttributeError as attr_error:
            return {"error": f"OpenAI API method not found: {str(attr_error)}"}
        except KeyError as key_error:
            return {"error": f"OpenAI API response format error: {str(key_error)}"}
        except Exception as openai_error:
            # If OpenAI fails, provide a basic fallback structure
            error_msg = str(openai_error)
            if "API key" in error_msg or "authentication" in error_msg.lower():
                return {"error": "OpenAI API authentication failed. Please check API key configuration."}
            else:
                return {"error": f"OpenAI API call failed: {str(openai_error)}"}

    except Exception as general_error:
        return {"error": f"Resume extraction failed: {str(general_error)}"}
    
    
def create_job_id():
    import random
    import string
    """Generate a simple unique job ID using a random base36 string."""
    chars = string.digits + string.ascii_lowercase
    job_id = "job_"



def fetch_job_by_id(job_id):
    """
    Fetch a job by its unique ID.
    Returns the job data if found, otherwise None.
    """
    job_data = fetch_job_data()
    for job in job_data:
        if job.get('id') == job_id:
            return job
    return None
def fetch_todays_activities():
    """
    Return a list of activities (candidate and job) that occurred today.
    Enhanced with new activity logger integration.
    """
    todays_activities = []
    today = datetime.datetime.now().date()
    
    # Try to get today's activities from new activity logger first
    try:
        from activity_logger import activity_logger
        new_todays_activities = activity_logger.get_todays_activities()
        
        # Convert new activity format to expected format
        for activity in new_todays_activities:
            todays_activities.append({
                'date': activity.get('timestamp', activity.get('date', '')),
                'description': activity.get('description', ''),
                'user': activity.get('user', 'Unknown'),
                'type': activity.get('activity_type', 'unknown'),
                'entity_type': activity.get('entity_type', ''),
                'entity_id': activity.get('entity_id', '')
            })
        
        print(f"✅ Fetched {len(new_todays_activities)} today's activities from new logger")
    except Exception as e:
        print(f"⚠️ Could not fetch today's activities from new activity logger: {e}")
    
    # Fallback to legacy method
    db_folder = os.path.join(os.path.dirname(__file__), 'db')
    candidate_file = os.path.join(db_folder, 'candidates.json')
    job_file = os.path.join(db_folder, 'jobs.json')
    
    # Get username mapping
    usernames = {}
    user_file = os.path.join(db_folder, 'userdata.json')
    if os.path.exists(user_file):
        with open(user_file, 'r') as f:
            try:
                users = json.load(f)
            except json.JSONDecodeError:
                users = []
        for u in users:
            if u.get('email'):
                usernames[u['email']] = u.get('username', '')
            if u.get('username'):
                usernames[u['username']] = u.get('username', '')
    
    # Legacy candidate activities for today
    if os.path.exists(candidate_file):
        with open(candidate_file, 'r') as f:
            try:
                candidates = json.load(f)
            except json.JSONDecodeError:
                candidates = []
        
        for c in candidates:
            # Check applied date
            if c.get('applied_date'):
                try:
                    applied_dt = datetime.datetime.strptime(c.get('applied_date'), '%Y-%m-%d')
                except Exception:
                    applied_dt = None
                if applied_dt and applied_dt.date() == today:
                    user_val = c.get('email', '')
                    user_val = usernames.get(user_val, user_val)
                    todays_activities.append({
                        'date': c.get('applied_date'),
                        'description': f"Candidate {c.get('name', '')} (ID: {c.get('id', '')}) applied",
                        'user': user_val if user_val else 'Unknown',
                        'type': 'candidate_applied',
                        'entity_type': 'candidate',
                        'entity_id': c.get('id', '')
                    })
            
            # Check updates today
            if c.get('updated_at') and c.get('updated_by'):
                try:
                    updated_dt = datetime.datetime.strptime(c.get('updated_at'), '%Y-%m-%d %H:%M:%S')
                except Exception:
                    updated_dt = None
                if updated_dt and updated_dt.date() == today:
                    user_val = c.get('updated_by', c.get('email', ''))
                    user_val = usernames.get(user_val, user_val)
                    candidate_name = c.get('name', '')
                    candidate_id = c.get('id', '')
                    
                    # Different types of updates
                    if c.get('status'):
                        todays_activities.append({
                            'date': c.get('updated_at'),
                            'description': f"Candidate {candidate_name} (ID: {candidate_id}) status changed to {c.get('status', '')}",
                            'user': user_val if user_val else 'Unknown',
                            'type': 'candidate_status_updated',
                            'entity_type': 'candidate',
                            'entity_id': candidate_id
                        })
                    
                    if c.get('interview_date') and c.get('interview_time'):
                        todays_activities.append({
                            'date': c.get('updated_at'),
                            'description': f"Interview scheduled for {candidate_name} (ID: {candidate_id}) on {c.get('interview_date', '')} at {c.get('interview_time', '')}",
                            'user': user_val if user_val else 'Unknown',
                            'type': 'interview_scheduled',
                            'entity_type': 'interview',
                            'entity_id': candidate_id
                        })
                    
                    if c.get('onboarding'):
                        todays_activities.append({
                            'date': c.get('updated_at'),
                            'description': f"Onboarding updated for {candidate_name} (ID: {candidate_id})",
                            'user': user_val if user_val else 'Unknown',
                            'type': 'onboarding_updated',
                            'entity_type': 'onboarding',
                            'entity_id': candidate_id
                        })
    
    # Legacy job activities for today
    if os.path.exists(job_file):
        with open(job_file, 'r') as f:
            try:
                jobs = json.load(f)
            except json.JSONDecodeError:
                jobs = []
        
        for job in jobs:
            posted_at = job.get('posted_at', '')
            try:
                posted_dt = datetime.datetime.strptime(posted_at, '%Y-%m-%d')
            except Exception:
                posted_dt = None
            if posted_dt and posted_dt.date() == today:
                user_val = job.get('job_posted_by', '')
                user_val = usernames.get(user_val, user_val)
                todays_activities.append({
                    'date': posted_at,
                    'description': f"Job posted: {job.get('job_title', '')}",
                    'user': user_val,
                    'type': 'job_posted',
                    'entity_type': 'job',
                    'entity_id': job.get('id', '')
                })
    
    # Remove duplicates
    unique_activities = []
    seen = set()
    for activity in todays_activities:
        key = f"{activity.get('date', '')}_{activity.get('description', '')}"
        if key not in seen:
            seen.add(key)
            unique_activities.append(activity)
    
    # Sort by date descending
    try:
        unique_activities = sorted(unique_activities, key=lambda x: x.get('date', ''), reverse=True)
    except Exception as e:
        print(f"⚠️ Error sorting today's activities: {e}")
    
    print(f"✅ Returning {len(unique_activities)} today's activities")
    return unique_activities